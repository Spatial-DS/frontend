import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def run_shelf_calculator(
    label_path, 
    dataset_path, 
    collection_mix_path, 
    output_path="shelfrun_final.docx",
    target_size=None,
    current_size=None,
    months=12
):
    # --- 1. PARAMETERS ---
    if target_size is not None and current_size is not None and float(current_size) != 0:
        constant = float(target_size) / float(current_size)
    else:
        constant = 0.75 
    no_of_months = float(months)

    print(f"Processing with: Constant={constant}, Months={no_of_months}")

    # --- 2. LOAD DATASETS ---
    try:
        with pd.ExcelFile(label_path) as xls:
            sheet_names = xls.sheet_names
            filter_sheet = next((s for s in sheet_names if s.lower() in ['filters', 'filiters']), None)
            if not filter_sheet: filter_sheet = sheet_names[0]
            shelf_sheet = next((s for s in sheet_names if s.lower() in ['shelf run', 'shelfrun']), None)
            if not shelf_sheet: shelf_sheet = sheet_names[1] if len(sheet_names) > 1 else sheet_names[0]
            print(f"Using Sheets -> Filters: '{filter_sheet}', Shelf Run: '{shelf_sheet}'")

            label_df = pd.read_excel(xls, dtype=str, sheet_name=filter_sheet)
            shelf_df = pd.read_excel(xls, dtype=str, sheet_name=shelf_sheet)

        data_df = pd.read_excel(dataset_path)
        # Renamed variable to reflect file content
        collection_mix_df = pd.read_excel(collection_mix_path)

    except Exception as e:
        print(f"Error loading excel files: {e}")
        return

    # --- 3. LOGIC ---
    collection_mix_df['Target end state collection'] = pd.to_numeric(collection_mix_df['Target end state collection'], errors='coerce').fillna(0)
    collection_mix_df['Retained'] = pd.to_numeric(collection_mix_df['Retained'], errors='coerce').fillna(0)
    collection_mix_df['holdings'] = collection_mix_df['Target end state collection'] + collection_mix_df['Retained']

    def set_to_float(df_col):
        return pd.to_numeric(df_col, errors='coerce').fillna(0)

    if 'Item DDC Class' in data_df.columns:
        data_df['Item DDC Class'] = data_df['Item DDC Class'].apply(lambda x: '000' if str(x).strip() == '0' else x)
    if 'Item DDC Class' in label_df.columns:
        label_df['Item DDC Class'] = label_df['Item DDC Class'].apply(lambda x: '000' if str(x).strip() == '0' else x)

    shelf_df['Avg Vol per m'] = set_to_float(shelf_df['Avg Vol per m'])
    shelf_df['No. of Tiers'] = set_to_float(shelf_df['No. of Tiers'])
    data_df['Loans and Renewals'] = set_to_float(data_df['Loans and Renewals'])
    data_df['Returns'] = set_to_float(data_df['Returns'])
    data_df['Potential RTOB Required'] = set_to_float(data_df['Potential RTOB Required'])
    collection_mix_df['holdings'] = set_to_float(collection_mix_df['holdings'])

    def create_cat(df):
        df['Cat'] = df['Category'].astype(str).str.strip() + "*" + df['Sub category'].astype(str).str.strip()

    create_cat(label_df)
    create_cat(shelf_df)
    create_cat(collection_mix_df)

    def pass_filiter(row, filt):
        for col, filt_val in filt.items():
            if col in ["Item Language","Item Age Lvl","Item Fiction Tag","Item Subject Suffix","Item DDC Class","Item Collection Code"]:
                filt_val = str(filt_val).strip()
                data_val = str(row.get(col, '')).strip() 
                if filt_val == '' or filt_val.lower() == 'nan': continue
                elif filt_val.startswith('NOT IN'):
                    exclude = [x.strip() for x in filt_val[7:].split(',')]
                    if data_val in exclude: return False
                elif ' or ' in filt_val:
                    options = [x.strip() for x in filt_val.split(' or ')]
                    if data_val not in options: return False
                else:
                    if data_val != filt_val: return False
        return True

    def assign_cat(row, filters):
        for f in filters:
            if pass_filiter(row, f): return f['Cat']
        return None

    label_df = label_df.fillna('').astype(str)
    for col in label_df.columns:
        if label_df[col].dtype == object: label_df[col] = label_df[col].str.strip()

    filters = label_df.to_dict(orient='records')
    print("Assigning categories...")
    data_df['Cat'] = data_df.apply(lambda row: assign_cat(row, filters), axis=1)

    collated = data_df.groupby('Cat', as_index=False).agg({
        'Loans and Renewals': 'sum',
        'Returns': 'sum',
        'Potential RTOB Required': 'sum'
    })

    finaldf = collated.merge(shelf_df, on='Cat', how='left')
    finaldf = finaldf.merge(collection_mix_df, on='Cat', how='left')

    def split_collection(df, original_cat, spine_suffix, front_suffix):
        row = df.loc[df['Cat'] == original_cat]
        if not row.empty:
            spine = row.copy()
            spine[['Loans and Renewals','Returns','Potential RTOB Required','holdings']] *= 0.8
            spine['Cat'] = original_cat.replace(original_cat.split('*')[1], spine_suffix)
            front = row.copy()
            front[['Loans and Renewals','Returns','Potential RTOB Required','holdings']] *= 0.2
            front['Cat'] = original_cat.replace(original_cat.split('*')[1], front_suffix)
            df = df.drop(row.index)
            df = pd.concat([df, spine, front], ignore_index=True)
        return df

    finaldf = split_collection(finaldf, "Children's Early Literacy Collection*Early Literacy (Baby) - English", "Early Literacy (Baby) - English (Spine)", "Early Literacy (Baby) - English (Front)")
    finaldf = split_collection(finaldf, "Children's Early Literacy Collection*Early Literacy (Baby) - Languages", "Early Literacy (Baby) - Languages (Spine)", "Early Literacy (Baby) - Languages (Front)")

    finaldf["Vol on shelf"]=finaldf['holdings']-(((finaldf['Loans and Renewals'])*constant)/no_of_months)+(((finaldf['Returns']+finaldf['Potential RTOB Required'])*constant)/no_of_months)
    finaldf["meter run"] = finaldf["Vol on shelf"] / finaldf["Avg Vol per m"]
    finaldf['meter run']=finaldf['meter run'].clip(lower=0.01)

    finaldf["Shelf run"] = finaldf["meter run"] / finaldf["No. of Tiers"]

    def fix_apostrophes(s):
        if pd.isna(s): return s
        return str(s).replace("’", "'").replace("‘", "'")

    finaldf['Cat'] = finaldf['Cat'].apply(fix_apostrophes)
    finaldf[['Category', 'Sub category']] = finaldf['Cat'].str.split('*', n=1, expand=True)

    # --- 4. GENERATE WORD DOC ---
    doc = Document()
    doc.add_heading('Shelf Run Calculation Results', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Date: {pd.Timestamp.now().strftime('%Y-%m-%d')}\n").bold = True
    p.add_run(f"Parameters: Target={target_size}, Current={current_size}, Constant={constant:.4f}, Months={no_of_months}")
    doc.add_paragraph().add_run().add_break() 

    df_shelfrun = finaldf[['Category','Sub category','meter run','No. of Tiers','Shelf run']]
    grouped = df_shelfrun.groupby("Category", sort=False)

    for category, group in grouped:
        if pd.isna(category): continue
        doc.add_heading(str(category), level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Sub category', 'Meter Run', 'No. of Tiers', 'Shelf Run']
        for i, text in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = text
            shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            cell.paragraphs[0].runs[0].bold = True

        for _, row in group.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Sub category'])
            m_run = row['meter run']
            row_cells[1].text = f"{m_run:.2f}" if pd.notnull(m_run) else "0.00"
            tiers = row['No. of Tiers']
            row_cells[2].text = str(int(tiers)) if pd.notnull(tiers) else "0"
            s_run = row['Shelf run']
            row_cells[3].text = f"{s_run:.1f}" if pd.notnull(s_run) else "0.0"
        doc.add_paragraph() 

    doc.save(output_path)
    print(f"Generated: {output_path}")